"""設計書 Markdown → Word (.docx) 一括変換スクリプト.

handson-lab/sample-app/docs/ の Markdown 設計書を
handson-lab/sample-app/docs/word/ に Word 形式で出力する。
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── 定数 ──
DOCS_DIR = Path(__file__).parent / "docs"
WORD_DIR = DOCS_DIR / "word"

# ドキュメントごとのメタ情報
DOC_META = {
    "system-overview.md": {
        "title": "システム概要設計書",
        "doc_id": "DOC-SYS-001",
        "version": "2.1",
    },
    "api-specification.md": {
        "title": "API 仕様書",
        "doc_id": "DOC-API-001",
        "version": "2.0",
    },
    "database-design.md": {
        "title": "データベース設計書",
        "doc_id": "DOC-DB-001",
        "version": "1.5",
    },
    "infrastructure.md": {
        "title": "インフラ構成設計書",
        "doc_id": "DOC-INFRA-001",
        "version": "1.3",
    },
    "security-design.md": {
        "title": "セキュリティ設計書",
        "doc_id": "DOC-SEC-001",
        "version": "1.4",
    },
}


def set_cell_shading(cell, color_hex: str):
    """テーブルセルの背景色を設定."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elem)


def apply_base_style(doc: Document):
    """ドキュメントの基本スタイルを設定."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "游ゴシック"
    font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # 見出しスタイル
    for level in range(1, 5):
        heading = doc.styles[f"Heading {level}"]
        heading.font.name = "游ゴシック"
        heading.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        heading.paragraph_format.space_before = Pt(18 - level * 2)
        heading.paragraph_format.space_after = Pt(8)
        if level == 1:
            heading.font.size = Pt(18)
        elif level == 2:
            heading.font.size = Pt(15)
        elif level == 3:
            heading.font.size = Pt(13)
        else:
            heading.font.size = Pt(11)


def add_cover_page(doc: Document, meta: dict):
    """表紙ページを追加."""
    for _ in range(6):
        doc.add_paragraph("")

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("社内業務ポータルシステム (BizPortal)")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run.bold = True

    doc.add_paragraph("")

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(meta["title"])
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run.bold = True

    for _ in range(4):
        doc.add_paragraph("")

    # メタ情報テーブル
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("ドキュメント ID", meta["doc_id"]),
        ("バージョン", meta["version"]),
        ("システム名", "社内業務ポータル (BizPortal)"),
        ("機密区分", "社外秘"),
    ]
    for i, (key, val) in enumerate(meta_rows):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = val
        for cell in table.row_cells(i):
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(11)

    doc.add_page_break()


def parse_markdown_line(line: str):
    """Markdown 行を解析して種別を返す."""
    stripped = line.rstrip()

    # 見出し
    heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
    if heading_match:
        level = len(heading_match.group(1))
        return ("heading", level, heading_match.group(2))

    # テーブル区切り行 (|---|---|)
    if re.match(r"^\|[\s\-:]+\|", stripped):
        return ("table_separator", None, None)

    # テーブル行
    if stripped.startswith("|") and stripped.endswith("|"):
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        return ("table_row", cells, None)

    # コードブロック開始/終了
    if stripped.startswith("```"):
        lang = stripped[3:].strip()
        return ("code_fence", lang, None)

    # 水平線
    if stripped in ("---", "***", "___"):
        return ("hr", None, None)

    # リスト項目
    list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", stripped)
    if list_match:
        indent = len(list_match.group(1))
        return ("list", indent, list_match.group(3))

    # 空行
    if not stripped:
        return ("blank", None, None)

    # 通常テキスト
    return ("text", None, stripped)


def render_inline(paragraph, text: str):
    """インラインの Markdown 書式 (bold, code, link) をレンダリング."""
    # 簡易的にボールド、インラインコード、リンクを処理
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)


def convert_md_to_docx(md_path: Path, meta: dict, output_path: Path):
    """Markdown ファイルを Word ドキュメントに変換."""
    doc = Document()
    apply_base_style(doc)

    # ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 表紙
    add_cover_page(doc, meta)

    # Markdown を読み込み
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    table_rows = []
    is_header_row = True

    i = 0
    while i < len(lines):
        line = lines[i]
        kind, val1, val2 = parse_markdown_line(line)

        # コードブロック
        if kind == "code_fence":
            if in_code_block:
                # コードブロック終了 → 整形して出力
                code_text = "\n".join(code_lines)
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # テーブル処理
        if kind == "table_row":
            table_rows.append(val1)
            i += 1
            # 次行がセパレータなら読み飛ばし
            if i < len(lines):
                next_kind, _, _ = parse_markdown_line(lines[i])
                if next_kind == "table_separator":
                    i += 1
            continue
        elif kind == "table_separator":
            i += 1
            continue
        else:
            # テーブル行以外に来たらテーブルを出力
            if table_rows:
                _flush_table(doc, table_rows)
                table_rows = []

        # 見出し
        if kind == "heading":
            level = val1
            doc.add_heading(val2, level=min(level, 4))

        # 水平線
        elif kind == "hr":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            run = para.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
            run.font.size = Pt(8)

        # リスト
        elif kind == "list":
            indent_level = val1 // 2
            para = doc.add_paragraph(style="List Bullet")
            render_inline(para, val2)
            para.paragraph_format.left_indent = Cm(1.5 + indent_level * 0.7)

        # テキスト
        elif kind == "text":
            para = doc.add_paragraph()
            render_inline(para, val2)

        # 空行はスキップ
        i += 1

    # 残りのテーブル
    if table_rows:
        _flush_table(doc, table_rows)

    doc.save(str(output_path))
    print(f"  ✅ {output_path.name}")


def _flush_table(doc: Document, rows: list[list[str]]):
    """蓄積したテーブル行を Word テーブルとして出力."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, cells in enumerate(rows):
        for col_idx, cell_text in enumerate(cells):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = ""
                para = cell.paragraphs[0]
                render_inline(para, cell_text)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9)

        # ヘッダー行のスタイル
        if row_idx == 0:
            for col_idx in range(num_cols):
                cell = table.cell(row_idx, col_idx)
                set_cell_shading(cell, "1A3C6E")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True


def main():
    """全設計書を Word に変換."""
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    print("📄 設計書 Markdown → Word 変換を開始します...\n")

    for md_file, meta in DOC_META.items():
        md_path = DOCS_DIR / md_file
        if not md_path.exists():
            print(f"  ⚠️ {md_file} が見つかりません。スキップします。")
            continue

        output_name = md_path.stem + ".docx"
        output_path = WORD_DIR / output_name
        convert_md_to_docx(md_path, meta, output_path)

    print(f"\n✅ 完了！Word ファイルは {WORD_DIR} に出力されました。")


if __name__ == "__main__":
    main()
